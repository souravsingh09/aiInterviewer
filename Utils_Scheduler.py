# This file contains all the auxiliary functions used in the main file.

import json
import os
import openai
import docx
import random
import string
import PyPDF2 as pypdf

from email import encoders
from email.mime.base import MIMEBase
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import ssl

import win32com.client
import pythoncom

from langchain.output_parsers.openai_functions import JsonOutputFunctionsParser
from langchain.prompts import ChatPromptTemplate
from langchain.chat_models import AzureChatOpenAI
from typing import List
from pydantic import BaseModel, Field
from langchain.utils.openai_functions import convert_pydantic_to_openai_function

from dotenv import load_dotenv

load_dotenv("credentials.env")

openai.api_key = os.environ["AZURE_OPENAI_API_KEY"]
openai.api_base = os.environ["AZURE_OPENAI_ENDPOINT"]
openai.api_type = 'azure'
openai.api_version = "2023-07-01-preview"

os.environ["OPENAI_API_TYPE"] = "azure"
os.environ["OPENAI_API_VERSION"] = "2023-07-01-preview"
os.environ["OPENAI_API_BASE"] = os.environ["AZURE_OPENAI_ENDPOINT"]
os.environ["OPENAI_API_KEY"] = os.environ["AZURE_OPENAI_API_KEY"]

class JDExtraction(BaseModel):
    """Tag the piece of text with particular info."""
    job_role: str = Field(description="Job role for which interview is being conducted.")
    vital_topics: List[str] = Field(description="Topics which are vital for the job role.")
    good_to_know_topics: List[str] = Field(description="Topics which are good to know for the job role.")

jd_function = [convert_pydantic_to_openai_function(JDExtraction)]

class ResumeExtraction(BaseModel):
    """Extract Resume related information"""
    candidate_email: str = Field(description="Email of the candidate")
    candidate_name: str = Field(description="Name of the candidate as mentioned in the resume")
    work_experience: List[str] = Field(description="The places where the candidate has worked")
    technical_skills: List[str] = Field(description="List of technical skills mentioned in the resume")
    project_titles: List[str] = Field(description="List of project names given in the resume")

resume_function = [convert_pydantic_to_openai_function(ResumeExtraction)]
model_jd = AzureChatOpenAI(temperature=0, deployment_name="gpt35turbo").bind(functions=jd_function)

model_resume  = AzureChatOpenAI(temperature=0, deployment_name="gpt35turbo").bind(functions=resume_function)

model_with_resume_function = model_resume.bind(
    functions=resume_function,
    function_call={"name": "ResumeExtraction"}
)

model_with_jd_functions = model_jd.bind(
    functions=jd_function,
    function_call={"name": "JDExtraction"}
)

class Utils_Scheduler:

    def __init__(self):
        self.info = {}

    def generate_key(self):
        alphanumeric_chars = string.ascii_letters + string.digits
        key = ''.join(random.choice(alphanumeric_chars) for _ in range(16))
        return key

    def job_description(self, file):

        # check if the file is pdf, docx or txt
        if file.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(file)
            # obtain the text from the pdf
            text = ""
            for page in pdf_reader.pages:
                text += page.extractText()
        elif file.endswith(".docx"):
            doc = docx.Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text
        elif file.endswith(".txt"):
            with open(file, "r") as f:
                text = f.read()
        
        # extract the job post, vital topics and good to know topics to ask questions about from the text using openai
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Think carefully, and then tag the text as instructed"),
            ("user", "{input}")
        ])

        tagging_chain = prompt | model_with_jd_functions | JsonOutputFunctionsParser() 
        result = tagging_chain.invoke({"input": text}) 
        print(result) 

        self.info["job_role"] = result["job_role"]
        self.info["vital_topics"] = result["vital_topics"]
        self.info["good_to_know_topics"] = result["good_to_know_topics"]

    def resume_extraction(self, file):

        # check if the file is pdf, docx or txt
        if file.endswith(".pdf"):
            pdf_reader = pypdf.PdfReader(file)
            # obtain the text from the pdf
            text = ""
            for page in pdf_reader.pages:
                text += page.extractText()
        elif file.endswith(".docx"):
            doc = docx.Document(file)
            text = ""
            for para in doc.paragraphs:
                text += para.text
        elif file.endswith(".txt"):
            with open(file, "r") as f:
                text = f.read()

        prompt = ChatPromptTemplate.from_messages([
            ("system", "Think carefully, and then tag the text as instructed"),
            ("user", "{input}")
        ])

        tagging_chain = prompt | model_with_resume_function | JsonOutputFunctionsParser() 
        result = tagging_chain.invoke({"input": text}) 
        print(result) 

        self.info["candidate_email"] = result["candidate_email"]
        self.info["candidate_name"] = result["candidate_name"]
        self.info["work_experience"] = result["work_experience"]
        self.info["technical_skills"] = result["technical_skills"]
        self.info["project_titles"] = result["project_titles"]

    def create_system_message(self):

        print(self.info)
        system_message =f"""

    ###CONTEXT###
    You are a very fastidious and to the point interviewer with good experience as a {self.info['job_role']}. You are conducting an oral interview.
    Today, you are interviewing a candidate for a {self.info['job_role']} position.
    The job description highlights the following vital topics: {",".join(self.info['vital_topics'])}.
    Additionally, it mentions some good-to-know topics: {",".join(self.info['good_to_know_topics'])}.
    Key information extracted from the candidate's resume is provided below:

    Candidate Name: {self.info['candidate_name']}
    Relevant Work Experience: {self.info['work_experience']}
    Technical Skills: {self.info['technical_skills']}
    Project Titles: {self.info['project_titles']}

    ###INSTRUCTIONS###
    Conduct a to the point interview with the candidate, focusing on the job requirements. Ask questions about each vital topic, each good-to-know topic, and the candidate's work experience and projects relevant for the job role. Allow the candidate to respond to each question before moving on to the next.
    Remember that the candidate may not always provide accurate or truthful answers or not answer some portions of your questions. In such cases, push back and ask clarification or a revised response.

    In a single dialog with the candidate always be to the point and only ask questions and respond in a neutral diplomatic tone without letting the candidate know whether they are correct or incorrect.

    Based on the candidate's answers for each topic:
    1. If the response is partially correct, just say "Okay. Please elaborate" regarding the topic to be elaborated.
    2. If the response is incorrect, contradictory, unsatisfactory or illogical, just say "Okay" and quickly move on to the next question on the same topic.
    3. If the response is correct, just say "Okay" and ask atleast 3, in-depth technical and conceptual, follow-up questions one question at a time in you dialog based on the candidate's answer. These questions can cover definitions, formulas, technologies, and tools used.
    4. If the candidate is unaware about any topic, quickly move on to the next topic.

    Before beginning the interview, introduce yourself as an "Clarissa" and welcome the candidate by their name which is already provided to you. Ask the candidate to introduce themselves, providing details such as their total years of experience, current company, and projects worked on. Based on that always start the interview focusing on the projects.

    When you conclude the interview, thank the candidate for the interview and inform that someone will communicate with the candidtate regarding this interview.

    """
        self.info["system_message"] = system_message


    def create_json(self, unique_id, resume_path, job_description_path, photograph_path):
        """
        This method will execute the following methos:
            job_description
            resume_extraction
            create_system_message
        Then it will save into a json file the following information:
            unique_id
            candidate_name
            system_message
            email to send the transcript and result to
        """
        self.job_description(job_description_path)
        self.resume_extraction(resume_path)
        self.create_system_message()

        data = {
            "unique_id": unique_id,
            "resume_path": resume_path,
            "job_description_path": job_description_path,
            "photo_path": photograph_path,
            "candidate_email": self.info["candidate_email"],
            "candidate_name": self.info["candidate_name"],
            "job_role": self.info["job_role"],
            "system_message": self.info["system_message"],
            "vital_topics": self.info["vital_topics"],
            "good_to_know_topics": self.info["good_to_know_topics"],
            "work_experience": self.info["work_experience"],
            "technical_skills": self.info["technical_skills"],
            "project_titles": self.info["project_titles"]

        }
        # save the data into a json file
        with open(f'static/scheduled_interviews/{unique_id}/data.json', 'w') as f:
            json.dump(data, f)

    def read_json(self, unique_id):
        """
        This method will read the json file created by the create_json method
        """
        with open(f'static/scheduled_interviews/{unique_id}/data.json', 'r') as f:
            self.info = json.load(f)
        return self.info
    
    def email_interview_details(self, email, url, unique_id):
        # Set up the SMTP server
        smtp_server = "smtp-mail.outlook.com"
        smtp_port = 587
        sender_email = ""
        sender_password = ""

        info = self.read_json(unique_id)
        recipients = [ email, info["candidate_email"]]

        message = MIMEMultipart()
        message["From"] = sender_email
        message["To"] = ", ".join(recipients)
        message["Subject"] = f"Interview Schedule and Link"
        body = f"""Hi {info["candidate_name"]},

You have been shortlisted for the interview for the job role of {info["job_role"]}. 
Kindly suggest a suitable time for the interview by replying to this email along with providing in attachment your most recent passport size photograph.


Also find in attachment the job description and your resume for your confirmation.

Regards,
Clarissa 
"""
        
        message.attach(MIMEText(body, "plain"))

        # Open the file in bynary mode
        file_paths = [info["job_description_path"], info["resume_path"], info["photo_path"]]
        for filename in file_paths:
            with open(filename, "rb") as attachment:
                # Add file as application/octet-stream
                # Email client can usually download this automatically as attachment
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())

            # Encode file in ASCII characters to send by email    
            encoders.encode_base64(part)

            # Add header as key/value pair to attachment part
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {filename}",
            )

            # Add attachment to message and convert message to string
            message.attach(part)
        text = message.as_string()

        # Log in to server using secure context and send email
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls(context=context)
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipients, text)
    
    def email_intimation(self, email, unique_id):
        outlook = win32com.client.Dispatch('Outlook.Application', pythoncom.CoInitialize())
        # size of the new email
        olmailitem = 0x0
        mail = outlook.CreateItem(olmailitem)
        info = self.read_json(unique_id)

        mail.To = info["candidate_email"]
        mail.Subject = "Shortlisted for Interview at Tech Mahindra"
        mail.Body = f"""Hi {info["candidate_name"]},

You have been shortlisted for the interview for the job role of {info["job_role"]}. 
Kindly suggest a suitable time for the interview by replying to this email, along with providing in attachment your most recent passport size photograph.


Also find in attachment the job description and your resume for your confirmation.

Regards,
Clarissa 
"""
        print(info["job_description_path"], info["resume_path"])
        mail.attachments.Add(info["job_description_path"])
        mail.attachments.Add(info["resume_path"])
        mail.CC = email

        mail.Send()

        return "Email sent successfully"


if __name__ == "__main__":
    Utils_Scheduler("Hello").recognize_from_microphone()